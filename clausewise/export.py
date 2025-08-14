from typing import List, Dict, Optional
from io import BytesIO
from docx import Document
from docx.shared import Pt


def _new_document() -> Document:
	doc = Document()
	style = doc.styles["Normal"]
	style.font.name = "Calibri"
	style.font.size = Pt(11)
	return doc


def generate_docx_bytes(text: str, title: Optional[str] = None) -> bytes:
	"""Create a simple DOCX file containing the given text and return its bytes."""
	doc = _new_document()
	if title:
		doc.add_heading(title, level=1)
	for para in text.split("\n\n"):
		p = doc.add_paragraph()
		p.add_run(para.strip())
	bio = BytesIO()
	doc.save(bio)
	return bio.getvalue()


def generate_docx_from_clauses(clauses: List[Dict], title: Optional[str] = None) -> bytes:
	"""Create a DOCX where each clause is a heading + paragraph."""
	doc = _new_document()
	if title:
		doc.add_heading(title, level=1)
	for clause in clauses:
		head = clause.get("title") or f"Clause {clause.get('id', '')}".strip()
		if head:
			doc.add_heading(head, level=2)
		p = doc.add_paragraph()
		p.add_run(clause.get("text", "").strip())
	bio = BytesIO()
	doc.save(bio)
	return bio.getvalue()
